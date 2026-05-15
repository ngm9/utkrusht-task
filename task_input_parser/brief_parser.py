"""Brief AST — deterministic structured representation of a customer brief.

The parser branches on file extension:
  .docx  → python-docx + zipfile (for embedded media)
  .md    → markdown-it-py
  .txt   → plain read with heuristic section splits

The agent operates on this AST, NOT on the raw file bytes. All entries are
Pydantic models so downstream consumers get validated, serialisable data.
"""
from __future__ import annotations

import io
import re
import zipfile
from hashlib import sha256
from pathlib import Path
from typing import List, Literal, Optional

from pydantic import BaseModel, Field


class Table(BaseModel):
    section_index: int
    headers: List[str]
    rows: List[List[str]]


class EmbeddedImage(BaseModel):
    image_ref: str               # stable identifier, e.g. "image_0"
    media_path_in_zip: str       # e.g. "word/media/image1.png"
    content_hash: str            # sha256(image_bytes)
    extension: str               # ".png", ".jpg", etc.
    size_bytes: int


class ExternalLink(BaseModel):
    url: str
    anchor_text: Optional[str] = None
    section_index: Optional[int] = None


class CodeFence(BaseModel):
    section_index: int
    language: Optional[str] = None
    content: str


class Section(BaseModel):
    index: int
    level: int                   # heading depth (1 = top-level)
    heading: Optional[str] = None
    paragraphs: List[str] = Field(default_factory=list)


class BriefAST(BaseModel):
    source_path: str
    source_format: Literal["docx", "md", "txt", "pdf"]
    source_hash: str             # sha256(file_bytes), for cache keying
    sections: List[Section] = Field(default_factory=list)
    tables: List[Table] = Field(default_factory=list)
    embedded_images: List[EmbeddedImage] = Field(default_factory=list)
    external_links: List[ExternalLink] = Field(default_factory=list)
    code_fences: List[CodeFence] = Field(default_factory=list)


def parse(path: Path) -> BriefAST:
    """Public dispatcher. Branches on file extension."""
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"Brief not found: {path}")
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return _parse_docx(path)
    if suffix in (".md", ".markdown"):
        return _parse_markdown(path)
    if suffix == ".txt":
        return _parse_text(path)
    if suffix == ".pdf":
        return _parse_pdf(path)
    raise ValueError(
        f"Unsupported brief format: {suffix!r}. "
        f"Supported formats: .docx, .md, .txt, .pdf"
    )


def _hash_file(path: Path) -> str:
    """Return the SHA-256 hex digest of a file's raw bytes."""
    return sha256(path.read_bytes()).hexdigest()


# ─────────────────────────────────────────────────────────────────────────────
# .docx
# ─────────────────────────────────────────────────────────────────────────────

def _parse_docx(path: Path) -> BriefAST:
    """Parse a .docx brief into a BriefAST using python-docx; extracts sections, tables, hyperlinks, and embedded images."""
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(str(path))
    sections: List[Section] = []
    tables: List[Table] = []
    external_links: List[ExternalLink] = []

    current_section = Section(index=0, level=1, heading=None, paragraphs=[])
    sections.append(current_section)

    for para in doc.paragraphs:
        style_name = (para.style.name or "").lower() if para.style else ""
        if style_name.startswith("heading"):
            try:
                level = int(style_name.replace("heading", "").strip() or "1")
            except ValueError:
                level = 1
            current_section = Section(
                index=len(sections),
                level=level,
                heading=para.text.strip() or None,
                paragraphs=[],
            )
            sections.append(current_section)
        else:
            text = para.text.strip()
            if text:
                current_section.paragraphs.append(text)
        # Hyperlinks live inside the paragraph's XML; extract them.
        for hyperlink in para._element.findall(".//" + qn("w:hyperlink")):
            r_id = hyperlink.get(qn("r:id"))
            if r_id and r_id in para.part.rels:
                url = para.part.rels[r_id].target_ref
                anchor = "".join(t.text or "" for t in hyperlink.findall(".//" + qn("w:t")))
                external_links.append(ExternalLink(
                    url=url,
                    anchor_text=anchor.strip() or None,
                    section_index=current_section.index,
                ))

    for ti, table in enumerate(doc.tables):
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if not rows:
            continue
        headers, *body = rows
        tables.append(Table(section_index=current_section.index, headers=headers, rows=body))

    # Embedded media via zipfile.
    embedded_images: List[EmbeddedImage] = []
    with zipfile.ZipFile(path) as zf:
        for ni, name in enumerate(zf.namelist()):
            if not name.startswith("word/media/"):
                continue
            data = zf.read(name)
            embedded_images.append(EmbeddedImage(
                image_ref=f"image_{ni}",
                media_path_in_zip=name,
                content_hash=sha256(data).hexdigest(),
                extension=Path(name).suffix.lower(),
                size_bytes=len(data),
            ))

    return BriefAST(
        source_path=str(path),
        source_format="docx",
        source_hash=_hash_file(path),
        sections=sections,
        tables=tables,
        embedded_images=embedded_images,
        external_links=external_links,
        code_fences=[],
    )


# ─────────────────────────────────────────────────────────────────────────────
# .md
# ─────────────────────────────────────────────────────────────────────────────

_LINK_RE = re.compile(r"\[([^\]]*)\]\((https?://[^)\s]+)\)")
_BARE_URL_RE = re.compile(r"(?<!\()(https?://[^\s)\]]+)")


def _parse_markdown(path: Path) -> BriefAST:
    """Parse a .md brief into a BriefAST using markdown-it-py; extracts sections, tables, code fences, and external links."""
    from markdown_it import MarkdownIt

    md = MarkdownIt("commonmark", {"html": False}).enable("table")
    tokens = md.parse(path.read_text(encoding="utf-8"))

    sections: List[Section] = [Section(index=0, level=1, heading=None, paragraphs=[])]
    tables: List[Table] = []
    code_fences: List[CodeFence] = []
    external_links: List[ExternalLink] = []

    current = sections[0]

    def push_section(level: int, heading: str) -> None:
        nonlocal current
        current = Section(index=len(sections), level=level, heading=heading.strip() or None, paragraphs=[])
        sections.append(current)

    i = 0
    while i < len(tokens):
        t = tokens[i]
        if t.type == "heading_open":
            level = int(t.tag[1:])  # "h2" → 2
            inline = tokens[i + 1].content if i + 1 < len(tokens) else ""
            push_section(level, inline)
            i += 3
            continue
        if t.type == "paragraph_open":
            inline = tokens[i + 1].content if i + 1 < len(tokens) else ""
            if inline.strip():
                current.paragraphs.append(inline.strip())
            for anchor, url in _LINK_RE.findall(inline):
                external_links.append(ExternalLink(url=url, anchor_text=anchor or None, section_index=current.index))
            for url in _BARE_URL_RE.findall(inline):
                external_links.append(ExternalLink(url=url, section_index=current.index))
            i += 3
            continue
        if t.type == "fence":
            code_fences.append(CodeFence(section_index=current.index, language=t.info.strip() or None, content=t.content))
            i += 1
            continue
        if t.type == "table_open":
            headers, rows = _collect_md_table(tokens, i)
            tables.append(Table(section_index=current.index, headers=headers, rows=rows))
            while i < len(tokens) and tokens[i].type != "table_close":
                i += 1
            i += 1
            continue
        i += 1

    return BriefAST(
        source_path=str(path),
        source_format="md",
        source_hash=_hash_file(path),
        sections=sections,
        tables=tables,
        embedded_images=[],
        external_links=external_links,
        code_fences=code_fences,
    )


def _collect_md_table(tokens, start_idx: int):
    """Walk markdown-it tokens from a table_open position and return (headers, rows)."""
    headers: List[str] = []
    rows: List[List[str]] = []
    in_head = False
    in_body = False
    current_row: List[str] = []

    for j in range(start_idx, len(tokens)):
        t = tokens[j]
        if t.type == "thead_open":
            in_head = True
        elif t.type == "thead_close":
            in_head = False
        elif t.type == "tbody_open":
            in_body = True
        elif t.type == "tbody_close":
            in_body = False
            break
        elif t.type == "tr_open":
            current_row = []
        elif t.type == "tr_close":
            if in_head and not headers:
                headers = current_row
            elif in_body:
                rows.append(current_row)
        elif t.type == "inline":
            current_row.append(t.content.strip())

    return headers, rows


# ─────────────────────────────────────────────────────────────────────────────
# .txt
# ─────────────────────────────────────────────────────────────────────────────

_HEADING_HEURISTIC_RE = re.compile(
    r"^("
    r"[A-Z][A-Z0-9 _\-]{3,}|"       # all-caps lines of reasonable length
    r"(Task|Question|Q)\s*\d+[:.].*|"  # "Task 1:" / "Q3." / "Question 2:"
    r"#+\s+.+|"                     # markdown-style heading in a .txt
    r".+:$"                         # line ending with a colon
    r")\s*$"
)


def _parse_text(path: Path) -> BriefAST:
    """Parse a .txt brief into a BriefAST using heuristic heading detection (all-caps lines, 'Task N:' prefixes, lines ending in colon)."""
    sections: List[Section] = []
    current = Section(index=0, level=1, heading=None, paragraphs=[])
    sections.append(current)
    external_links: List[ExternalLink] = []
    buffer: List[str] = []

    def flush_paragraph():
        if buffer:
            para = " ".join(buffer).strip()
            if para:
                current.paragraphs.append(para)
                for anchor, url in _LINK_RE.findall(para):
                    external_links.append(ExternalLink(url=url, anchor_text=anchor or None, section_index=current.index))
                for url in _BARE_URL_RE.findall(para):
                    external_links.append(ExternalLink(url=url, section_index=current.index))
            buffer.clear()

    for raw_line in path.read_text(encoding="utf-8").splitlines():
        line = raw_line.rstrip()
        if not line:
            flush_paragraph()
            continue
        if _HEADING_HEURISTIC_RE.match(line):
            flush_paragraph()
            level = line.count("#") if line.startswith("#") else 1
            heading = line.lstrip("#").strip().rstrip(":")
            current = Section(index=len(sections), level=level, heading=heading or None, paragraphs=[])
            sections.append(current)
            continue
        buffer.append(line.strip())
    flush_paragraph()

    return BriefAST(
        source_path=str(path),
        source_format="txt",
        source_hash=_hash_file(path),
        sections=sections,
        tables=[],
        embedded_images=[],
        external_links=external_links,
        code_fences=[],
    )


# ─────────────────────────────────────────────────────────────────────────────
# .pdf
# ─────────────────────────────────────────────────────────────────────────────

def _parse_pdf(path: Path) -> BriefAST:
    """Parse a .pdf brief into a BriefAST using pdfplumber.

    Extracts text page-by-page, detects headings with the same heuristic used
    for .txt files, extracts tables from each page, and collects hyperlinks.
    Embedded images are not extracted (no Drive upload path for PDF media).
    """
    import pdfplumber

    sections: List[Section] = []
    tables: List[Table] = []
    external_links: List[ExternalLink] = []
    current = Section(index=0, level=1, heading=None, paragraphs=[])
    sections.append(current)
    buffer: List[str] = []

    def flush_paragraph() -> None:
        """Flush buffered lines into the current section as a single paragraph."""
        if buffer:
            para = " ".join(buffer).strip()
            if para:
                current.paragraphs.append(para)
                for anchor, url in _LINK_RE.findall(para):
                    external_links.append(ExternalLink(url=url, anchor_text=anchor or None, section_index=current.index))
                for url in _BARE_URL_RE.findall(para):
                    external_links.append(ExternalLink(url=url, section_index=current.index))
            buffer.clear()

    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            # ── Extract hyperlinks from page annotations ──────────────────────
            for annot in (page.annots or []):
                uri = annot.get("uri") or annot.get("URI") or ""
                if uri.startswith("http"):
                    external_links.append(ExternalLink(
                        url=uri,
                        section_index=current.index,
                    ))

            # ── Extract tables before text so table cells aren't double-counted
            for pdf_table in page.extract_tables() or []:
                if not pdf_table:
                    continue
                rows = [[cell.strip() if cell else "" for cell in row] for row in pdf_table]
                rows = [r for r in rows if any(r)]  # drop fully empty rows
                if len(rows) < 2:
                    continue
                headers, *body = rows
                tables.append(Table(
                    section_index=current.index,
                    headers=headers,
                    rows=body,
                ))

            # ── Extract text lines, detect headings heuristically ─────────────
            text = page.extract_text(x_tolerance=2, y_tolerance=2) or ""
            for raw_line in text.splitlines():
                line = raw_line.rstrip()
                if not line:
                    flush_paragraph()
                    continue
                if _HEADING_HEURISTIC_RE.match(line):
                    flush_paragraph()
                    level = line.count("#") if line.startswith("#") else 1
                    heading = line.lstrip("#").strip().rstrip(":")
                    current = Section(
                        index=len(sections),
                        level=level,
                        heading=heading or None,
                        paragraphs=[],
                    )
                    sections.append(current)
                    continue
                buffer.append(line.strip())
            flush_paragraph()

    return BriefAST(
        source_path=str(path),
        source_format="pdf",
        source_hash=_hash_file(path),
        sections=sections,
        tables=tables,
        embedded_images=[],
        external_links=external_links,
        code_fences=[],
    )
